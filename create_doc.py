from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create(employee_name,designation,date,base_income,pera,deductions,deduction_total):
    # employee_name = 'Jasper Nichol M. Fabella'
    # designation = 'Programmer III'
    # date = 'October 1-15 2019'
    # base_income = 9000
    # pera = 1000
    gross_income = base_income + pera
    #
    # deductions = [
    #     ['PHILHEALTH',1000]
    # ]


    if os.path.exists("{}".format('payslip.docx')):
        os.remove("{}".format('payslip.docx'))

    document = Document()
    document.add_heading('PHRMO SALARY SLIP FOR {}'.format(date), 1)

    # p = document.add_paragraph('\nSTAFF NAME:     ')
    # p.add_run('{}'.format(employee_name)).bold = True
    # p.add_run('\nDESIGNATION:     ')
    # p.add_run('{}'.format(designation)).bold = True
    # p.add_run('\nDATE:     ')
    # p.add_run('{}'.format(date)).bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True


    employee_details = (
        ('STAFF NAME:',employee_name),
        ('DESIGNATION:',designation),
        ('DATE:',date)
    )

    table = document.add_table(rows=1, cols=2)

    for qty, id in employee_details:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        #row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[1].text = id
        row_cells[1].paragraphs[0].runs[0].font.bold = True

    p = document.add_paragraph('', style='Intense Quote')
    p = document.add_paragraph('BASE SALARY:     ')
    p.add_run('{:,.2f}'.format(base_income)).bold = True
    p.add_run('\nPERA:     ')
    p.add_run('{:,.2f}'.format(pera)).bold = True
    p.add_run('\nGROSS INCOME:     ')
    p.add_run('{:,.2f}'.format(gross_income)).bold = True

    p = document.add_paragraph('', style='Intense Quote')
    p.paragraph_format.alignment= WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('This is a computer generated slip no signature required').font.size = Pt(8)

    #document.add_picture('monty-truth.png', width=Inches(1.25))

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[1].text = 'DEDUCTIONS'
    hdr_cells[1].paragraphs[0].runs[0].font.bold
    for qty, id in deductions:
        row_cells = table.add_row().cells
        row_cells[1].text = str(qty)
        row_cells[2].text = '{:,.2f}'.format(id)


    p = document.add_paragraph('', style='Intense Quote')
    p = document.add_paragraph()



    p = document.add_paragraph('NET INCOME:     ')
    p.add_run('{:,.2f}'.format(gross_income - deduction_total)).bold = True
    document.save('payslip.docx')
    os.system('start payslip.docx')